"""
08_grey_zone_distribution.py
=============================
Quantification of the distribution of FBI scores across three clinically
interpretable decision bands:

  - FBI < 9   (below screening cut-off — rule-out zone)
  - FBI 9-14  (intermediate / "grey zone")
  - FBI >= 15 (high-specificity confirmatory band — rule-in zone)

For each band the script reports:
  - Numerosity and proportion of the cohort
  - Proportion of consensus-impaired patients with Wilson 95% CI

Wilson CIs (Wilson 1927) are used instead of the normal-approximation
(Wald) interval because some bands — particularly FBI >= 15 — contain
relatively few patients, where the Wald interval performs poorly
(coverage below nominal) and can extend outside [0,1].

"""
from __future__ import annotations
from pathlib import Path
import sys
import numpy as np
import pandas as pd

sys.path.insert(0, str(Path(__file__).parent))
from helpers import wilson_ci

# ── configuration ─────────────────────────────────────────────────────────────
INPUT_SAV = Path("../data/cognitivita_FINAL_recoded_ETATEST_ECASFRSBEv3.sav")
INPUT_CSV = Path("../data/analytic_dataset.csv")
OUTDIR    = Path("../outputs")
OUTDIR.mkdir(parents=True, exist_ok=True)

WORD_OUTDIR = OUTDIR / "word_tables"
WORD_OUTDIR.mkdir(parents=True, exist_ok=True)


# ── data loading ──────────────────────────────────────────────────────────────
def load_data():
    """Returns (full FBI cohort, complete-case cohort)."""
    if INPUT_SAV.exists():
        import pyreadstat
        df, _  = pyreadstat.read_sav(str(INPUT_SAV))
        sla    = df[df["SLA_NONSLA"] == "SLA"].copy()
        full   = sla[sla["FBI_total_score"].notna()].copy()
        anch   = ["ECAS_BEH_PATOL", "FRSBE_TOTALE_PATOL", "BBI_PATOL"]
        cc     = full[full[anch].notna().all(axis=1)].copy()
        cc["gold"] = (cc[anch].astype(int).sum(axis=1) >= 2).astype(int)
        return full, cc
    elif INPUT_CSV.exists():
        df   = pd.read_csv(INPUT_CSV)
        df   = df.rename(columns={"fbi_total": "FBI_total_score",
                                   "gold_consensus": "gold"})
        full = df.copy()
        cc   = df[df["in_complete_case"] == 1].copy()
        return full, cc
    else:
        raise FileNotFoundError("Provide either the SAV or the analytic CSV.")


# ── Word export ───────────────────────────────────────────────────────────────
def save_word(df_out, title, legend, filename):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[Word export] python-docx not available — skipping.")
        return

    HEADER_FILL = "D9E2F3"; FONT = "Calibri"

    def shade(cell, col):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), col); tcPr.append(shd)

    def settext(cell, text, bold=False, center=False):
        cell.text = ""
        p   = cell.paragraphs[0]
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.bold = bold; run.font.size = Pt(9); run.font.name = FONT

    doc = Document()
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(10)
    section = doc.sections[0]
    for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
        setattr(section, attr, Inches(0.7))

    p = doc.add_paragraph()
    r = p.add_run(title); r.bold = True; r.font.size = Pt(11); r.font.name = FONT
    p.paragraph_format.space_after = Pt(6)

    tbl = doc.add_table(rows=1, cols=len(df_out.columns))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df_out.columns):
        settext(tbl.rows[0].cells[j], col, bold=True, center=True)
        shade(tbl.rows[0].cells[j], HEADER_FILL)
    for _, row in df_out.iterrows():
        cells = tbl.add_row().cells
        for j, col in enumerate(df_out.columns):
            settext(cells[j], row[col], center=(j > 0))

    if legend:
        pl = doc.add_paragraph()
        pl.paragraph_format.space_before = Pt(6)
        rl = pl.add_run(legend)
        rl.italic = True; rl.font.size = Pt(9); rl.font.name = FONT

    outpath = WORD_OUTDIR / filename
    doc.save(outpath)
    print(f"Written: {outpath}")


# ── main analysis ─────────────────────────────────────────────────────────────
def main():
    full, cc = load_data()

    print("=" * 78)
    print("DISTRIBUTION ACROSS DECISION BANDS (grey-zone quantification)")
    print("=" * 78)
    print("\nThree clinically interpretable score bands:")
    print("  FBI <  9  : below screening cut-off (rule-out)")
    print("  FBI 9-14  : intermediate / 'grey zone'")
    print("  FBI >= 15 : confirmatory band (rule-in)\n")

    # Band sizes in each cohort (no gold standard needed)
    for label, X in [("FBI cohort (n=" + str(len(full)) + ")", full),
                     ("Complete-case (n=" + str(len(cc)) + ")", cc)]:
        f = X["FBI_total_score"]
        n = len(X)
        nl = int((f < 9).sum())
        nm = int(((f >= 9) & (f <= 14)).sum())
        nh = int((f >= 15).sum())
        print(f"\n{label}:")
        print(f"  FBI <  9  : {nl:4d}  ({nl/n*100:5.1f}%)")
        print(f"  FBI 9-14  : {nm:4d}  ({nm/n*100:5.1f}%)")
        print(f"  FBI >= 15 : {nh:4d}  ({nh/n*100:5.1f}%)")

    # Impaired proportion within each band (complete-case only, Wilson CI)
    print("\n" + "-" * 78)
    print("Consensus-impaired proportion within each band  (complete-case, Wilson 95% CI)")
    print("-" * 78)

    bands = [
        (0,  8,  "FBI < 9"),
        (9,  11, "FBI 9-11"),
        (9,  14, "FBI 9-14 (grey zone)"),
        (12, 14, "FBI 12-14"),
        (15, 99, "FBI >= 15"),
    ]

    rows_csv  = []
    rows_word = []
    for lo, hi, lab in bands:
        m   = (cc["FBI_total_score"] >= lo) & (cc["FBI_total_score"] <= hi)
        band = cc[m]
        n   = len(band)
        imp = int((band["gold"] == 1).sum())
        if n == 0:
            continue
        pct        = imp / n * 100
        ci_lo, ci_hi = wilson_ci(imp, n)
        print(f"  {lab:22s}: n={n:3d}  impaired={imp:3d} "
              f"({pct:5.1f}%, 95% CI {ci_lo*100:.1f}%-{ci_hi*100:.1f}%)")
        rows_csv.append({"band": lab, "n": n, "impaired": imp,
                         "pct_impaired": pct,
                         "wilson_ci_lo": ci_lo * 100,
                         "wilson_ci_hi": ci_hi * 100})
        rows_word.append({
            "FBI score band"       : lab,
            "n"                    : str(n),
            "Impaired, n (%)"      : f"{imp} ({pct:.1f}%)",
            "Wilson 95% CI"        : f"{ci_lo*100:.1f}%\u2013{ci_hi*100:.1f}%",
        })

    csv_path = OUTDIR / "table8_grey_zone.csv"
    pd.DataFrame(rows_csv).to_csv(csv_path, index=False)
    print(f"\nSaved: {csv_path}")

    save_word(
        pd.DataFrame(rows_word),
        title="Table 8. Distribution of FBI scores across decision bands.",
        legend=(
            "Complete-case subset (n = 345). Consensus-impaired = flagged by "
            "\u22652 of 3 FBI-independent instruments (ECAS-behavioural, FrSBe, BBI). "
            "Wilson (1927) score 95% confidence intervals are used instead of the "
            "normal-approximation interval because some bands contain few patients."
        ),
        filename="Table8_grey_zone.docx",
    )


if __name__ == "__main__":
    main()