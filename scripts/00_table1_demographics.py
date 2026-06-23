"""
Build Table 1 (demographics) from the analytic dataset, plus Word (.docx)
versions of Table 1-5 for direct use in a manuscript.

"""
from pathlib import Path
import pandas as pd
import numpy as np
from scipy.stats import mannwhitneyu, chi2_contingency, fisher_exact

DATA   = Path("../data/analytic_dataset.csv")
OUTDIR = Path("../outputs")
OUT    = OUTDIR / "table1_demographics.csv"

df = pd.read_csv(DATA)
cc  = df[df['in_complete_case'] == 1].copy()   # n=345
ncc = df[df['in_complete_case'] == 0].copy()   # n=161  (non-complete-case)

# ── statistical test helpers ─────────────────────────────────────────────────

def fmt_p(p):
    if p is None or np.isnan(p):
        return "—"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"

def mw_p(col):
    """Mann-Whitney U p-value: complete-case vs non-complete-case."""
    a = cc[col].dropna()
    b = ncc[col].dropna()
    if len(a) < 3 or len(b) < 3:
        return None
    _, p = mannwhitneyu(a, b, alternative='two-sided')
    return p

def cat_p(col, vals):
    """Chi-square (or Fisher) p-value for a categorical variable."""
    counts_cc  = [(cc[col] == v).sum()  for v in vals]
    counts_ncc = [(ncc[col] == v).sum() for v in vals]
    table = np.array([counts_cc, counts_ncc])
    # drop columns that are zero in BOTH groups (truly unobserved category)
    keep = table.sum(axis=0) > 0
    table = table[:, keep]
    if table.shape[1] < 2:
        return None
    # Always attempt Fisher exact on 2x2; for larger tables try chi2 first
    # and fall back to collapsing into present/absent binary if chi2 fails
    # (e.g. expected cell = 0 after removing empty columns).
    if table.shape == (2, 2):
        _, p = fisher_exact(table)
    else:
        try:
            _, p, _, expected = chi2_contingency(table, correction=False)
            if (expected < 1).any() or (expected < 5).mean() > 0.2:
                raise ValueError("low expected counts")
        except (ValueError, Exception):
            # collapse to binary: any category vs none
            a = int((cc[col].isin(vals)).sum())
            b = int((ncc[col].isin(vals)).sum())
            t = np.array([[a, len(cc)-a],[b, len(ncc)-b]])
            _, p = fisher_exact(t)
    return p

# ── row builder ───────────────────────────────────────────────────────────────

rows = []

def add(name, full_val, cc_val, p_val=""):
    rows.append({
        "Characteristic"          : name,
        "FBI cohort (n=506)"      : full_val,
        "Complete-case (n=345)"   : cc_val,
        "p-value\u00b9"           : p_val,
    })

def descr(s):
    nn = s.notna().sum()
    if nn == 0: return "\u2014"
    return f"{s.mean():.1f} \u00b1 {s.std():.1f}"

def pct(s, val):
    return f"{(s==val).sum()} ({(s==val).mean()*100:.1f}%)"

def mitos_kings(s, levels):
    return " / ".join(str((s==L).sum()) for L in levels)

# Continuous
add("Age at test, years (mean \u00b1 SD)",
    descr(df['age_at_test']), descr(cc['age_at_test']),
    fmt_p(mw_p('age_at_test')))

add("Male sex, n (%)",
    pct(df['sex'], 'M'), pct(cc['sex'], 'M'),
    fmt_p(cat_p('sex', ['M'])))

add("Education, years (mean \u00b1 SD)",
    descr(df['education_years']), descr(cc['education_years']),
    fmt_p(mw_p('education_years')))

# Onset site
b1 = (df['onset_site']=='Bulbar').sum();  s1 = (df['onset_site']=='Spinal').sum()
b2 = (cc['onset_site']=='Bulbar').sum();  s2 = (cc['onset_site']=='Spinal').sum()
add("Site of onset (n with data)", f"{b1+s1}", f"{b2+s2}",
    fmt_p(cat_p('onset_site', ['Bulbar','Spinal'])))
add("   Bulbar, n (%)",
    f"{b1} ({b1/(b1+s1)*100:.1f}%)", f"{b2} ({b2/(b2+s2)*100:.1f}%)", "")
add("   Spinal, n (%)",
    f"{s1} ({s1/(b1+s1)*100:.1f}%)", f"{s2} ({s2/(b2+s2)*100:.1f}%)", "")

add("ALSFRS-R total at assessment (mean \u00b1 SD)",
    f"{descr(df['alsfrs_r_total'])} (n\u2009=\u2009{df['alsfrs_r_total'].notna().sum()})",
    f"{descr(cc['alsfrs_r_total'])} (n\u2009=\u2009{cc['alsfrs_r_total'].notna().sum()})",
    fmt_p(mw_p('alsfrs_r_total')))

# MiToS
add("MiToS stage (n with data)",
    f"{df['mitos'].notna().sum()}", f"{cc['mitos'].notna().sum()}",
    fmt_p(cat_p('mitos', [0,1,2])))
add("MiToS stage 0 / 1 / 2, n",
    mitos_kings(df['mitos'], [0,1,2]),
    mitos_kings(cc['mitos'], [0,1,2]), "")

# King's
add("King\u2019s stage (n with data)",
    f"{df['kings'].notna().sum()}", f"{cc['kings'].notna().sum()}",
    fmt_p(cat_p('kings', [1,2,3,4])))
add("King\u2019s stage 1 / 2 / 3 / 4, n",
    mitos_kings(df['kings'], [1,2,3,4]),
    mitos_kings(cc['kings'], [1,2,3,4]), "")

# C9orf72
c9p_full = (df['c9orf72']=='Positive').sum()
c9t_full = df['c9orf72'].isin(['Positive','Negative']).sum()
c9p_cc   = (cc['c9orf72']=='Positive').sum()
c9t_cc   = cc['c9orf72'].isin(['Positive','Negative']).sum()
add("C9orf72 expansion (positive / tested)",
    f"{c9p_full} / {c9t_full} ({c9p_full/c9t_full*100:.1f}%)",
    f"{c9p_cc} / {c9t_cc} ({c9p_cc/c9t_cc*100:.1f}%)",
    fmt_p(cat_p('c9orf72', ['Positive'])))

# Strong 2017
add("Strong 2017 classification (n with data)",
    f"{df['strong_2017'].notna().sum()}", f"{cc['strong_2017'].notna().sum()}",
    fmt_p(cat_p('strong_2017', ['CN','ALSci','ALSbi','ALScbi','FTD'])))
add("Strong 2017 classification, n", "", "", "")
for cat in ['CN','ALSci','ALSbi','ALScbi','FTD']:
    nice = {'CN':'   Cognitively normal','ALSci':'   ALSci','ALSbi':'   ALSbi',
            'ALScbi':'   ALScbi','FTD':'   ALS-FTD'}[cat]
    add(nice, str((df['strong_2017']==cat).sum()),
        str((cc['strong_2017']==cat).sum()), "")

# FBI score
add("FBI total (mean \u00b1 SD; median, range)",
    f"{descr(df['fbi_total'])}; {df['fbi_total'].median():.0f} "
    f"({df['fbi_total'].min():.0f}\u2013{df['fbi_total'].max():.0f})",
    f"{descr(cc['fbi_total'])}; {cc['fbi_total'].median():.0f} "
    f"({cc['fbi_total'].min():.0f}\u2013{cc['fbi_total'].max():.0f})",
    fmt_p(mw_p('fbi_total')))

add("Behavioural impairment by consensus, n (%)",
    "\u2014",
    f"{int(cc['gold_consensus'].sum())} ({cc['gold_consensus'].mean()*100:.1f}%)",
    "\u2014")

# ── save CSV ──────────────────────────────────────────────────────────────────
table1 = pd.DataFrame(rows)
table1.to_csv(OUT, index=False)
print(table1.to_string(index=False))
print(f"\nWritten: {OUT}")

# QA: Strong 2017 category sums
for label, df_use in [('FBI cohort', df), ('Complete-case', cc)]:
    strong_sum = sum((df_use['strong_2017']==c).sum()
                     for c in ['CN','ALSci','ALSbi','ALScbi','FTD'])
    strong_nn  = df_use['strong_2017'].notna().sum()
    if strong_sum != strong_nn:
        print(f"\nWARNING ({label}): Strong 2017 categories sum to {strong_sum} but "
              f"{strong_nn} patients have non-missing strong_2017 — "
              f"{strong_nn - strong_sum} patient(s) uncategorised. "
              f"Check source coding before publishing.")

# ============================================================================
# WORD TABLE EXPORT
# ============================================================================
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    print("\n[Word export] python-docx not installed — skipping. "
          "Install with:  pip install python-docx --break-system-packages")

WORD_OUTDIR = OUTDIR / "word_tables"
HEADER_FILL = "D9E2F3"
FONT_NAME   = "Calibri"

def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _set_cell_text(cell, text, bold=False, size=10, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold       = bold
    run.font.size  = Pt(size)
    run.font.name  = FONT_NAME

def dataframe_to_docx_table(doc, df_in, header_size=10, body_size=10,
                             center_from_col=1):
    """Insert df_in as a Word table. Columns >= center_from_col are centred."""
    table = doc.add_table(rows=1, cols=len(df_in.columns))
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df_in.columns):
        _set_cell_text(hdr_cells[j], col, bold=True, size=header_size,
                       align_center=(j >= center_from_col))
        _shade_cell(hdr_cells[j], HEADER_FILL)

    for _, row in df_in.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df_in.columns):
            _set_cell_text(cells[j], row[col], bold=False, size=body_size,
                           align_center=(j >= center_from_col))
    return table

def make_table_docx(df_in, title, legend, filename, landscape=None):
    if not _DOCX_AVAILABLE:
        return
    if landscape is None:
        landscape = len(df_in.columns) > 6

    doc     = Document()
    style   = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (section.page_height,
                                                    section.page_width)
    for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
        setattr(section, attr, Inches(0.7))

    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    p_title.paragraph_format.space_after = Pt(8)

    dataframe_to_docx_table(doc, df_in)

    if legend:
        p_leg = doc.add_paragraph()
        p_leg.paragraph_format.space_before = Pt(8)
        run = p_leg.add_run(legend)
        run.italic    = True
        run.font.size = Pt(9)
        run.font.name = FONT_NAME

    WORD_OUTDIR.mkdir(parents=True, exist_ok=True)
    outpath = WORD_OUTDIR / filename
    doc.save(outpath)
    print(f"Written: {outpath}")

# ── Table 1 Word export ───────────────────────────────────────────────────────
if _DOCX_AVAILABLE:
    make_table_docx(
        table1,
        title="Table 1. Demographic and clinical characteristics of the cohort.",
        legend=(
            "Continuous variables: mean \u00b1 SD; categorical variables: n (%). "
            "SD, standard deviation; ALSFRS-R, Amyotrophic Lateral Sclerosis "
            "Functional Rating Scale \u2013 Revised. Complete-case\u2009=\u2009patients "
            "with concurrent ECAS-behavioural, FrSBe and BBI data. "
            "\u00b9p-value compares complete-case (n\u2009=\u2009345) vs "
            "non-complete-case (n\u2009=\u2009161) patients: Mann\u2013Whitney U "
            "for continuous variables; chi-square or Fisher exact for categorical "
            "variables. p shown only for the summary row of grouped variables "
            "(e.g. onset site, MiToS, King\u2019s, Strong 2017); subcategory rows "
            "left blank to avoid redundancy. \u2014\u2009=\u2009not applicable."
        ),
        filename="Table1_demographics.docx",
    )

# ── Formatters for Tables 2-5 (unchanged) ────────────────────────────────────

def fmt_table2(path):
    d = pd.read_csv(path)
    out_rows = []
    for _, r in d.iterrows():
        label = (f"FBI \u2265 {int(r['cutoff'])}"
                 + (" (legacy)" if int(r['cutoff']) == 25 else ""))
        out_rows.append({
            'Cut-off'    : label,
            'Sensitivity': f"{r['sens']:.3f}",
            'Specificity': f"{r['spec']:.3f}",
            'PPV'        : f"{r['ppv']:.3f}",
            'NPV'        : f"{r['npv']:.3f}",
            "Cohen\u2019s \u03ba": f"{r['kappa']:+.3f}",
        })
    out = pd.DataFrame(out_rows)
    n_total    = int(d.loc[0, ['tp','fp','tn','fn']].sum())
    prevalence = (d.loc[0,'tp'] + d.loc[0,'fn']) / n_total
    legend = (
        f"Performance of candidate FBI total cut-offs against the consensus "
        f"reference standard (n\u2009=\u2009{n_total}, behavioural-impairment "
        f"prevalence {prevalence*100:.1f}%). PPV, positive predictive value; "
        f"NPV, negative predictive value. The proposed ALS cut-off is "
        f"Youden-optimal; FBI \u2265 25 is the legacy bvFTD-derived threshold, "
        f"shown for comparison."
    )
    return out, legend

def fmt_table3(path):
    d = pd.read_csv(path)
    out_rows = []
    for _, r in d.iterrows():
        out_rows.append({
            'Subscale' : r['subscale'],
            'Anchor'   : r['anchor'],
            'n'        : f"{int(r['n'])}",
            'AUC (95% CI)': (f"{r['auc']:.3f} "
                             f"({r['auc_lo']:.3f}\u2013{r['auc_hi']:.3f})"),
            'Optimal cut-off (95% CI)': (f"\u2265 {r['optimal_cutoff']:.0f} "
                                         f"({r['cutoff_lo']:.0f}\u2013"
                                         f"{r['cutoff_hi']:.0f})"),
            'Sens.\u2009/\u2009Spec.': (f"{r['sensitivity']:.2f} / "
                                         f"{r['specificity']:.2f}"),
            "Cohen\u2019s \u03ba": f"{r['kappa']:+.3f}",
        })
    out    = pd.DataFrame(out_rows)
    legend = (
        "Subscale recalibration. Youden-optimal cut-offs for the FBI "
        "Apathy/negative (items 1\u201312) and Disinhibition (items 13\u201324) "
        "subscales, against the multi-instrument consensus reference standard "
        "and against construct-matched ECAS/FrSBe subscales. AUC\u2009=\u2009area "
        "under the ROC curve. Bootstrap 95% CI from 2000 stratified resamples."
    )
    return out, legend

def fmt_table4(path):
    d = pd.read_csv(path)
    out_rows = []
    for _, r in d.iterrows():
        out_rows.append({
            'Reference standard'     : r['reference'],
            'n'                      : f"{int(r['n'])}",
            'Sens. at \u22659'       : f"{r['sens_at_9']:.3f}",
            '\u03ba at \u22659'      : f"{r['kappa_at_9']:+.3f}",
            '\u03ba at \u226525'     : f"{r['kappa_at_25']:+.3f}",
            '\u0394\u03ba'           : f"{r['delta_kappa']:+.3f}",
        })
    out      = pd.DataFrame(out_rows)
    n_values = sorted(d['n'].unique())
    if len(n_values) == 1:
        n_note = f"n\u2009=\u2009{int(n_values[0])} for all comparisons."
    else:
        n_note = (
            f"n varies by row ({', '.join(str(int(v)) for v in n_values)}) "
            f"because the LCA-7 reference standard requires four additional "
            f"indicators not available for the entire complete-case sample "
            f"\u2014 see the n column."
        )
    legend = (
        f"Cohen\u2019s \u03ba of the dichotomised FBI total against six "
        f"reference standards, at the proposed (\u22659) and legacy (\u226525) "
        f"cut-offs. \u0394\u03ba\u2009=\u2009improvement with the new cut-off. "
        f"{n_note}"
    )
    return out, legend

def fmt_table5(path):
    d = pd.read_csv(path)
    out_rows = []
    for _, r in d.iterrows():
        out_rows.append({
            'Subgroup'               : r['subgroup'],
            'n'                      : f"{int(r['n'])}",
            'Gold prevalence'        : f"{r['gold_prevalence']*100:.1f}%",
            'AUC'                    : f"{r['auc']:.3f}",
            'Youden cut-off (95% CI)': (f"\u2265 {r['youden_cutoff']:.0f} "
                                        f"({r['cutoff_lo']:.0f}\u2013"
                                        f"{r['cutoff_hi']:.0f})"),
            'Sens\u2009/\u2009Spec\u2009/\u2009\u03ba at \u22659' :
                (f"{r['sens_at_9']:.2f}\u2009/\u2009{r['spec_at_9']:.2f}"
                 f"\u2009/\u2009{r['kappa_at_9']:+.3f}"),
            'Sens\u2009/\u2009Spec\u2009/\u2009\u03ba at \u226525':
                (f"{r['sens_at_25']:.2f}\u2009/\u2009{r['spec_at_25']:.2f}"
                 f"\u2009/\u2009{r['kappa_at_25']:+.3f}"),
            'FBI mean\u2009\u00b1\u2009SD (median)':
                (f"{r['fbi_mean']:.2f}\u2009\u00b1\u2009{r['fbi_sd']:.2f} "
                 f"({r['fbi_median']:.0f})"),
        })
    out    = pd.DataFrame(out_rows)
    legend = (
        "Sensitivity analysis: FBI performance stratified by site of onset "
        "(bulbar vs spinal). Supplementary table \u2014 not one of the "
        "manuscript\u2019s four numbered tables; renumber/integrate as "
        "appropriate (e.g. as Supplementary Table S1) when assembling the "
        "submission."
    )
    return out, legend

# ── Tables 2-5 Word export ────────────────────────────────────────────────────
if _DOCX_AVAILABLE:
    for path, fmt_fn, title, fname in [
        (OUTDIR / "table2_cutoffs.csv",        fmt_table2,
         "Table 2. Performance of candidate FBI total cut-offs.",
         "Table2_cutoffs.docx"),
        (OUTDIR / "table3_subscales.csv",       fmt_table3,
         "Table 3. Subscale recalibration.",
         "Table3_subscales.docx"),
        (OUTDIR / "table4_kappa.csv",           fmt_table4,
         "Table 4. Agreement (Cohen\u2019s \u03ba) against six reference standards.",
         "Table4_kappa.docx"),
        (OUTDIR / "table5_bulbar_spinal.csv",   fmt_table5,
         "Table 5. Sensitivity analysis \u2014 bulbar vs spinal onset.",
         "Table5_bulbar_spinal.docx"),
    ]:
        if path.exists():
            df_fmt, legend = fmt_fn(path)
            make_table_docx(df_fmt, title=title, legend=legend, filename=fname)
        else:
            print(f"[Word export] {path} not found \u2014 run the script that "
                  f"produces it first (02/03/04). Skipping {fname}.")