"""
09_decision_curve_analysis.py
==============================
Decision Curve Analysis (Vickers & Elkin 2006, Med Decis Making 26:565-574)
for FBI cut-offs in ALS. Quantifies clinical utility of candidate cut-offs
by estimating net benefit (NB) across the clinically relevant range of
threshold probabilities (pt).

Net benefit:
    NB(pt) = TP/N - FP/N * (pt / (1 - pt))

Strategies compared:
    1. FBI continuous  (logistic regression FBI -> P(impaired), quadratic)
    2. FBI >= 9        (proposed cut-off)
    3. FBI >= 12       (confirmatory band)
    4. FBI >= 25       (legacy cut-off)
    5. Treat all       (baseline: assume everyone impaired)
    6. Treat none      (NB = 0 by definition)

Output:
    - figures/FigureS1_DCA.png  (300 dpi)
    - figures/FigureS1_DCA.pdf
    - outputs/log_dca.txt
    - outputs/table9_dca_nb.csv  + Word table
    Tabulated NB at pt in {0.10, 0.20, 0.25, 0.30, 0.40, 0.50}

IMPORTANT NOTE ON THE CONTINUOUS MODEL
----------------------------------------
The quadratic logistic model is estimated and evaluated on the SAME
complete-case sample (n = 345). Its net-benefit estimates are therefore
optimistic relative to the fixed cut-off strategies (FBI >= 9 / 12 / 25),
which have no free parameters estimated from data and are thus immune to
this form of in-sample overfitting. At threshold probabilities close to the
observed prevalence (~0.25), where the continuous model most benefits from
fitting, this optimism can make it appear artificially superior to FBI >= 9.
The plot and table flag this limitation explicitly. In a prospective external
validation the continuous model would be expected to partially regress toward
the performance of the simple cut-off rules.

"""
from __future__ import annotations
from pathlib import Path
import sys
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import statsmodels.api as sm

sys.path.insert(0, str(Path(__file__).parent))

# ── configuration ─────────────────────────────────────────────────────────────
SEED = 2026
np.random.seed(SEED)

INPUT_SAV = Path("../data/cognitivita_FINAL_recoded_ETATEST_ECASFRSBEv3.sav")
INPUT_CSV = Path("../data/analytic_dataset.csv")
OUTDIR    = Path("../outputs")
FIGDIR    = Path("../figures")
OUTDIR.mkdir(parents=True, exist_ok=True)
FIGDIR.mkdir(parents=True, exist_ok=True)

WORD_OUTDIR = OUTDIR / "word_tables"
WORD_OUTDIR.mkdir(parents=True, exist_ok=True)


# ── data loading ──────────────────────────────────────────────────────────────
def load_data():
    if INPUT_SAV.exists():
        import pyreadstat
        df, _ = pyreadstat.read_sav(str(INPUT_SAV))
        sla   = df[df["SLA_NONSLA"] == "SLA"].copy()
        df    = sla[sla["FBI_total_score"].notna()].copy()
        anch  = ["ECAS_BEH_PATOL", "FRSBE_TOTALE_PATOL", "BBI_PATOL"]
        cc    = df[df[anch].notna().all(axis=1)].copy()
        cc["gold"] = (cc[anch].astype(int).sum(axis=1) >= 2).astype(int)
        return cc[["FBI_total_score", "gold"]].copy()
    elif INPUT_CSV.exists():
        cc = pd.read_csv(INPUT_CSV)
        cc = cc[cc["in_complete_case"] == 1].copy()
        cc = cc.rename(columns={"fbi_total": "FBI_total_score",
                                 "gold_consensus": "gold"})
        return cc[["FBI_total_score", "gold"]].copy()
    else:
        raise FileNotFoundError("Provide either the SAV or the analytic CSV.")


# ── net benefit functions ─────────────────────────────────────────────────────
def nb_treat_all(y, pt):
    p = y.mean()
    return p - (1 - p) * (pt / (1 - pt))

def nb_predictor(y, p_pred, pt):
    pos = (p_pred >= pt)
    tp  = (pos & (y == 1)).sum()
    fp  = (pos & (y == 0)).sum()
    return tp / len(y) - (fp / len(y)) * (pt / (1 - pt))

def nb_binary(y, classifier, pt):
    tp = ((classifier == 1) & (y == 1)).sum()
    fp = ((classifier == 1) & (y == 0)).sum()
    return tp / len(y) - (fp / len(y)) * (pt / (1 - pt))


# ── Word export ───────────────────────────────────────────────────────────────
def save_word(df_out, title, legend, filename):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[Word export] python-docx not available — skipping.")
        return

    HEADER_FILL = "D9E2F3"; FONT = "Calibri"

    def shade(cell, col):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), col); tcPr.append(shd)

    def settext(cell, text, bold=False, center=False):
        cell.text = ""
        p   = cell.paragraphs[0]
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.bold = bold; run.font.size = Pt(9); run.font.name = FONT

    doc = Document()
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(10)
    section = doc.sections[0]
    for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
        setattr(section, attr, Inches(0.7))

    p = doc.add_paragraph()
    r = p.add_run(title); r.bold = True; r.font.size = Pt(11); r.font.name = FONT
    p.paragraph_format.space_after = Pt(6)

    tbl = doc.add_table(rows=1, cols=len(df_out.columns))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df_out.columns):
        settext(tbl.rows[0].cells[j], col, bold=True, center=True)
        shade(tbl.rows[0].cells[j], HEADER_FILL)
    for _, row in df_out.iterrows():
        cells = tbl.add_row().cells
        for j, col in enumerate(df_out.columns):
            settext(cells[j], row[col], center=(j > 0))

    if legend:
        pl = doc.add_paragraph()
        pl.paragraph_format.space_before = Pt(6)
        rl = pl.add_run(legend)
        rl.italic = True; rl.font.size = Pt(9); rl.font.name = FONT

    outpath = WORD_OUTDIR / filename
    doc.save(outpath)
    print(f"Written: {outpath}")


# ── main analysis ─────────────────────────────────────────────────────────────
def main():
    cc   = load_data()
    y    = cc["gold"].astype(int).values
    fbi  = cc["FBI_total_score"].astype(float).values
    N    = len(y)
    prev = y.mean()

    log = open(OUTDIR / "log_dca.txt", "w", encoding="utf-8")
    def say(*a):
        msg = " ".join(str(x) for x in a)
        print(msg); log.write(msg + "\n"); log.flush()

    say("=" * 72)
    say("DECISION CURVE ANALYSIS - FBI in ALS")
    say("=" * 72)
    say(f"\nN = {N}, prevalence = {prev:.4f} ({y.sum()}/{N})")
    say("\nNOTE ON THE CONTINUOUS MODEL: the quadratic logistic model is")
    say("estimated and evaluated on the SAME sample. Its NB estimates are")
    say("therefore optimistic vs the fixed cut-off strategies. See module")
    say("docstring for details.")

    # Continuous FBI via quadratic logistic regression
    X   = sm.add_constant(np.column_stack([fbi, fbi**2]))
    mod = sm.Logit(y, X).fit(disp=0)
    p_continuous = mod.predict(X)
    say(f"\nLogistic model (quadratic FBI): log-likelihood = {mod.llf:.2f}")
    for fval in [0, 9, 25]:
        say(f"  P(impaired | FBI = {fval:2d}) = "
            f"{mod.predict(np.array([[1, fval, fval**2]]))[0]:.3f}")

    # Threshold-probability grid
    pt_grid = np.linspace(0.01, 0.60, 200)
    nb_all  = np.array([nb_treat_all(y, pt) for pt in pt_grid])
    nb_cont = np.array([nb_predictor(y, p_continuous, pt) for pt in pt_grid])
    cls_9   = (fbi >= 9).astype(int)
    cls_12  = (fbi >= 12).astype(int)
    cls_25  = (fbi >= 25).astype(int)
    nb_9    = np.array([nb_binary(y, cls_9,  pt) for pt in pt_grid])
    nb_12   = np.array([nb_binary(y, cls_12, pt) for pt in pt_grid])
    nb_25   = np.array([nb_binary(y, cls_25, pt) for pt in pt_grid])

    # Tabulate NB at representative threshold probabilities
    pt_check = [0.10, 0.20, 0.25, 0.30, 0.40, 0.50]
    strategies = [
        ("Continuous FBI*", nb_cont),
        ("FBI >= 9",         nb_9),
        ("FBI >= 12",        nb_12),
        ("FBI >= 25",        nb_25),
        ("Treat all",        nb_all),
    ]
    say("\n" + "-" * 72)
    say("NET BENEFIT at representative threshold probabilities")
    say("* Continuous FBI NB is optimistic (in-sample model, see NOTE above)")
    say("-" * 72)
    hdr = f"{'Strategy':22s}" + "".join(f"  pt={pt:.2f}" for pt in pt_check)
    say(hdr); say("-" * len(hdr))
    rows_csv  = []
    rows_word = []
    for name, arr in strategies:
        vals = []
        for pt_chk in pt_check:
            idx = np.argmin(np.abs(pt_grid - pt_chk))
            vals.append(f"{arr[idx]:+.4f}")
        say(f"{name:22s}" + "".join(f"  {v:>8s}" for v in vals))
        rows_csv.append({"Strategy": name, **{f"pt={pt:.2f}": v
                         for pt, v in zip(pt_check, vals)}})
        rows_word.append({"Strategy": name, **{f"pt = {pt:.2f}": v
                          for pt, v in zip(pt_check, vals)}})

    csv_path = OUTDIR / "table9_dca_nb.csv"
    pd.DataFrame(rows_csv).to_csv(csv_path, index=False)
    say(f"\nSaved: {csv_path}")

    # Interventions avoided per 100 at pt = 0.25
    say("\n" + "=" * 72)
    say("INTERVENTIONS AVOIDED PER 100 PATIENTS vs treat-all  (pt = 0.25)")
    say("=" * 72)
    pt_clin  = 0.25
    idx_clin = np.argmin(np.abs(pt_grid - pt_clin))
    nb_all_at = nb_all[idx_clin]
    for name, arr in strategies[:-1]:   # skip treat-all
        delta  = arr[idx_clin] - nb_all_at
        avoided = delta * 100 / (pt_clin / (1 - pt_clin))
        say(f"  {name:22s}: NB={arr[idx_clin]:+.4f}, "
            f"delta vs treat-all={delta:+.4f}, "
            f"interventions avoided/100={avoided:+.1f}")

    # Word table
    save_word(
        pd.DataFrame(rows_word),
        title="Table 9. Net benefit at representative threshold probabilities (Decision Curve Analysis).",
        legend=(
            "Net benefit NB(pt) = TP/N \u2212 (FP/N) \u00d7 pt/(1\u2212pt). "
            "Higher NB = greater clinical utility. "
            "Treat-none NB is always 0 (not shown). "
            "* Continuous FBI NB is optimistic: the logistic model is estimated "
            "and evaluated on the same sample (n\u2009=\u2009345), whereas the fixed "
            "cut-off strategies have no free parameters and are thus unaffected "
            "by in-sample overfitting."
        ),
        filename="Table9_dca_nb.docx",
    )

    # ── FIGURE S1 ──────────────────────────────────────────────────────────
    plt.rcParams.update({"font.size": 11, "font.family": "DejaVu Sans",
                          "axes.linewidth": 1.1})
    fig, ax = plt.subplots(figsize=(8.5, 6.5))

    ax.plot(pt_grid, nb_all,  lw=1.6, color="#888888", ls="--", label="Treat all")
    ax.axhline(0,             lw=1.6, color="#444444", ls=":",  label="Treat none (NB = 0)")
    ax.plot(pt_grid, nb_cont, lw=2.8, color="#1f4e79",
            label="FBI continuous* (optimistic — in-sample model)")
    ax.plot(pt_grid, nb_9,    lw=2.4, color="#2e7d32",
            label="FBI \u2265 9 (proposed cut-off)")
    ax.plot(pt_grid, nb_12,   lw=2.0, color="#6a994e", ls="-.",
            label="FBI \u2265 12 (confirmatory band)")
    ax.plot(pt_grid, nb_25,   lw=2.4, color="#ed7d31",
            label="FBI \u2265 25 (legacy cut-off)")

    ax.axvspan(0.10, 0.40, alpha=0.07, color="#1f4e79", zorder=0)
    ax.text(0.25, 0.005, "Clinical decision range\n(pt 10\u201340%)",
            ha="center", va="bottom", fontsize=9.5, color="#1f4e79",
            style="italic", alpha=0.85)
    ax.axvline(prev, color="black", lw=0.9, alpha=0.4, ls=":")
    ax.text(prev + 0.005, 0.22, f"Cohort\nprevalence\n({prev:.0%})",
            fontsize=9, color="black", alpha=0.6, va="center")

    # Add a note about the continuous model directly on the plot
    ax.text(0.58, nb_cont[np.argmin(np.abs(pt_grid - 0.58))] + 0.012,
            "* in-sample,\noptimistic",
            fontsize=8, color="#1f4e79", ha="right", va="bottom", style="italic")

    ax.set_xlabel("Threshold probability of behavioural impairment, $p_t$",
                  fontsize=11.5)
    ax.set_ylabel("Net benefit", fontsize=11.5)
    ax.set_title(
        "Decision Curve Analysis — clinical utility of FBI cut-offs in ALS\n"
        "* Continuous model NB is optimistic (estimated and evaluated in-sample)",
        fontsize=11.5, fontweight="bold")
    ax.set_xlim(0, 0.60)
    # Extend lower bound slightly below 0 to show strategies worse than treat-none
    y_min = min(nb_25.min(), nb_12.min(), -0.01)
    ax.set_ylim(y_min - 0.005, 0.27)
    ax.grid(alpha=0.25)
    ax.legend(loc="upper right", framealpha=0.95, fontsize=9.5)

    plt.tight_layout()
    plt.savefig(FIGDIR / "FigureS1_DCA.png", dpi=300, bbox_inches="tight")
    plt.savefig(FIGDIR / "FigureS1_DCA.pdf", bbox_inches="tight")
    plt.close()
    say(f"\nSaved: {FIGDIR / 'FigureS1_DCA.png'}")
    say(f"Saved: {FIGDIR / 'FigureS1_DCA.pdf'}")
    log.close()


if __name__ == "__main__":
    main()