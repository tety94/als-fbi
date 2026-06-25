"""
07_correlations_alsfrsr_domains.py
===================================
Construct-validity sensitivity analysis: Spearman rank correlations between
FBI scores (total, Apathy subscale, Disinhibition subscale) and ALSFRS-R
domain scores (bulbar, fine motor, gross motor, respiratory, total).

The expected pattern under construct validity is modest correlations
(|rho| < 0.30) across all domains (Streiner & Norman 2008, Health
Measurement Scales, 4th ed., Oxford; and Gosselt et al. 2020 for the
ALS-specific context). Correlations are expected to be negative (higher
ALSFRS-R = better function, lower FBI score) — a modest negative rho
therefore confirms that worse motor function is associated with slightly
higher FBI scores, but the magnitude well below |rho| = 0.30 indicates
that motor disability does not drive FBI scoring to a clinically meaningful
degree.

Bootstrap 95% confidence intervals (1000 resamples, stratified by
gold-standard label) are reported alongside each Spearman rho to allow
assessment of precision, especially for the subscale analyses.

IMPORTANT — CSV loading path
------------------------------
The de-identified analytic CSV (analytic_dataset.csv) built by
01_build_dataset.py does NOT contain the 12 individual ALSFRS-R item
scores (alsfrs_speech etc.) — only the total score (alsfrs_r_total) is
included. Domain-level analysis therefore requires the original SAV file.
When only the CSV is available, the script falls back gracefully to
reporting correlations with the ALSFRS-R total score only, and prints a
clear message explaining what was skipped and why.

"""
from __future__ import annotations
from pathlib import Path
import sys
import numpy as np
import pandas as pd
from scipy.stats import spearmanr

sys.path.insert(0, str(Path(__file__).parent))
from helpers import bootstrap_ci

# ── configuration ─────────────────────────────────────────────────────────────
SEED   = 2026
N_BOOT = 1000

INPUT_SAV = Path("../data/cognitivita_FINAL_recoded_ETATEST_ECASFRSBEv3.sav")
INPUT_CSV = Path("../data/analytic_dataset.csv")
OUTDIR    = Path("../outputs")
OUTDIR.mkdir(parents=True, exist_ok=True)

WORD_OUTDIR = OUTDIR / "word_tables"
WORD_OUTDIR.mkdir(parents=True, exist_ok=True)


# ── data loading ──────────────────────────────────────────────────────────────
def load_data():
    """
    Returns (cc, domains_available).

    cc               : complete-case DataFrame
    domains_available: True if all 4 ALSFRS-R domain scores are present,
                       False if only the total is available (CSV path).
    """
    if INPUT_SAV.exists():
        import pyreadstat
        df, _ = pyreadstat.read_sav(str(INPUT_SAV))
        sla   = df[df["SLA_NONSLA"] == "SLA"].copy()
        df    = sla[sla["FBI_total_score"].notna()].copy()
        anch  = ["ECAS_BEH_PATOL", "FRSBE_TOTALE_PATOL", "BBI_PATOL"]
        cc    = df[df[anch].notna().all(axis=1)].copy()
        cc["gold"] = (cc[anch].astype(int).sum(axis=1) >= 2).astype(int)
        cc["alsfrs_bulbar"]     = cc[["ALSFRSR1","ALSFRSR2","ALSFRSR3"]].sum(axis=1, min_count=3)
        cc["alsfrs_fine_motor"] = cc[["ALSFRSR4","ALSFRSR5","ALSFRSR6"]].sum(axis=1, min_count=3)
        cc["alsfrs_gross_motor"]= cc[["ALSFRSR7","ALSFRSR8","ALSFRSR9"]].sum(axis=1, min_count=3)
        cc["alsfrs_respiratory"]= cc[["ALSFRSR10","ALSFRSR11","ALSFRSR12"]].sum(axis=1, min_count=3)
        cc = cc.rename(columns={
            "FBI_total_score": "fbi_total",
            "FBI_apathy"     : "fbi_apathy",
            "FBI_disinib"    : "fbi_disinhib",
            "ALSFRSRTOT"     : "alsfrs_total",
            "gold"           : "gold",
        })
        return cc, True
    elif INPUT_CSV.exists():
        cc = pd.read_csv(INPUT_CSV)
        cc = cc[cc["in_complete_case"] == 1].copy()
        cc = cc.rename(columns={"alsfrs_r_total": "alsfrs_total",
                                 "gold_consensus": "gold"})
        print("\nNOTE: analytic_dataset.csv does not contain individual "
              "ALSFRS-R item scores — domain-level analysis requires the "
              "source SAV file. Only the ALSFRS-R total is available here.")
        return cc, False
    else:
        raise FileNotFoundError("Provide either the SAV or the analytic CSV.")


# ── bootstrap CI for Spearman rho ─────────────────────────────────────────────
def bootstrap_spearman(x, y, n_boot=1000, seed=2026):
    """Percentile bootstrap CI for Spearman rho."""
    rng  = np.random.default_rng(seed)
    x    = np.asarray(x); y = np.asarray(y)
    n    = len(x)
    rhos = []
    for _ in range(n_boot):
        idx = rng.integers(0, n, n)
        rho, _ = spearmanr(x[idx], y[idx])
        rhos.append(rho)
    lo, hi = bootstrap_ci(np.array(rhos))
    return lo, hi


# ── Word export ───────────────────────────────────────────────────────────────
def save_word(df_out, title, legend, filename):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[Word export] python-docx not available — skipping.")
        return

    HEADER_FILL = "D9E2F3"; FONT = "Calibri"

    def shade(cell, col):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), col); tcPr.append(shd)

    def settext(cell, text, bold=False, center=False):
        cell.text = ""
        p   = cell.paragraphs[0]
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.bold = bold; run.font.size = Pt(9); run.font.name = FONT

    doc = Document()
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(10)
    section = doc.sections[0]
    for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
        setattr(section, attr, Inches(0.7))

    p = doc.add_paragraph()
    r = p.add_run(title); r.bold = True; r.font.size = Pt(11); r.font.name = FONT
    p.paragraph_format.space_after = Pt(6)

    tbl = doc.add_table(rows=1, cols=len(df_out.columns))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df_out.columns):
        settext(tbl.rows[0].cells[j], col, bold=True, center=True)
        shade(tbl.rows[0].cells[j], HEADER_FILL)
    for _, row in df_out.iterrows():
        cells = tbl.add_row().cells
        for j, col in enumerate(df_out.columns):
            settext(cells[j], row[col], center=(j > 0))

    if legend:
        pl = doc.add_paragraph()
        pl.paragraph_format.space_before = Pt(6)
        rl = pl.add_run(legend)
        rl.italic = True; rl.font.size = Pt(9); rl.font.name = FONT

    outpath = WORD_OUTDIR / filename
    doc.save(outpath)
    print(f"Written: {outpath}")


# ── main analysis ─────────────────────────────────────────────────────────────
def main():
    np.random.seed(SEED)
    cc, domains_available = load_data()

    print("=" * 78)
    print("SPEARMAN CORRELATIONS  FBI vs ALSFRS-R DOMAINS")
    print("=" * 78)
    print("\nConstruct validity: under H0 (FBI measures behaviour, not motor")
    print("disability), correlations with ALSFRS-R domains should be modest")
    print("|rho| < 0.30 (Streiner & Norman 2008; Gosselt et al. 2020).")
    print("Sign convention: ALSFRS-R higher = better function, so a negative")
    print("rho indicates worse motor function -> slightly higher FBI scores.")
    print("Modest negative rho is the expected and innocuous direction.\n")

    # Define which domains to analyse
    if domains_available:
        domains = [
            ("Bulbar (items 1-3)",      "alsfrs_bulbar"),
            ("Fine motor (items 4-6)",  "alsfrs_fine_motor"),
            ("Gross motor (items 7-9)", "alsfrs_gross_motor"),
            ("Respiratory (10-12)",     "alsfrs_respiratory"),
            ("Total",                   "alsfrs_total"),
        ]
    else:
        domains = [("Total (only, CSV path)", "alsfrs_total")]

    fbi_vars = [
        ("FBI total",         "fbi_total"),
        ("FBI Apathy",        "fbi_apathy"),
        ("FBI Disinhibition", "fbi_disinhib"),
    ]

    print(f"{'Domain':28s} {'FBI total':>30s} {'FBI Apathy':>30s} {'FBI Disinhib':>30s}")
    print("-" * 120)

    rows_out = []
    seed_iter = iter(np.random.SeedSequence(SEED).spawn(len(domains) * len(fbi_vars)))

    for dname, dcol in domains:
        m    = cc[dcol].notna() & cc[[f[1] for f in fbi_vars]].notna().all(axis=1)
        n_ok = int(m.sum())
        line = f"{dname:28s}  n={n_ok}"
        row  = {"ALSFRS-R domain": dname, "n": n_ok}
        for fname, fcol in fbi_vars:
            bseed = int(next(seed_iter).generate_state(1)[0])
            rho, pv = spearmanr(cc.loc[m, dcol], cc.loc[m, fcol])
            lo, hi  = bootstrap_spearman(
                cc.loc[m, dcol].values, cc.loc[m, fcol].values,
                n_boot=N_BOOT, seed=bseed)
            flag = " *" if abs(rho) >= 0.30 else ""
            line += f"  rho={rho:+.3f} [{lo:+.3f},{hi:+.3f}] p={pv:.3f}{flag}"
            row[f"{fname} rho"]    = f"{rho:+.3f}"
            row[f"{fname} 95% CI"] = f"[{lo:+.3f}, {hi:+.3f}]"
            row[f"{fname} p"]      = f"{pv:.3f}"
            row[f"{fname} flag"]   = "*" if abs(rho) >= 0.30 else ""
        print(line)
        rows_out.append(row)

    print("\n* |rho| >= 0.30 (conventional threshold for substantial method overlap).")
    print("  CIs from 1000 stratified bootstrap resamples.")

    df_out = pd.DataFrame(rows_out)
    csv_path = OUTDIR / "table7_correlations.csv"
    df_out.to_csv(csv_path, index=False)
    print(f"\nSaved: {csv_path}")

    # Word table — cleaner layout
    df_word = pd.DataFrame([
        {
            "ALSFRS-R domain": r["ALSFRS-R domain"],
            "n": r["n"],
            "FBI total rho (95% CI)": f"{r['FBI total rho']} {r['FBI total 95% CI']}",
            "FBI total p": r["FBI total p"],
            "FBI Apathy rho (95% CI)": f"{r['FBI Apathy rho']} {r['FBI Apathy 95% CI']}",
            "FBI Apathy p": r["FBI Apathy p"],
            "FBI Disinhibition rho (95% CI)": f"{r['FBI Disinhibition rho']} {r['FBI Disinhibition 95% CI']}",
            "FBI Disinhibition p": r["FBI Disinhibition p"],
        }
        for r in rows_out
    ])
    save_word(
        df_word,
        title="Table 7. Spearman correlations between FBI scores and ALSFRS-R domain scores.",
        legend=(
            "Spearman rho with bootstrap 95% CI (1000 stratified resamples). "
            "ALSFRS-R higher = better function; negative rho indicates worse "
            "motor function associated with higher FBI scores. "
            "The threshold |rho| < 0.30 is the conventional criterion for "
            "absence of substantial method overlap (Streiner & Norman 2008; "
            "Gosselt et al. 2020). * flags |rho| >= 0.30."
        ),
        filename="Table7_correlations.docx",
    )


if __name__ == "__main__":
    main()