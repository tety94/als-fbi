"""
06_sensitivity_motor_severity.py
=================================
Sensitivity analysis: ROC recalibration of the FBI total cut-off stratified
by motor disability severity (ALSFRS-R total > 35 vs <= 35).

Reproduces the "Sensitivity by motor severity" section in the manuscript.

Output (printed and saved as CSV):
  - AUC, Youden cut-off and bootstrap 95% CI for each subgroup
  - Sensitivity/specificity/kappa at FBI >= 9 and FBI >= 25 in each subgroup
  - Mann-Whitney comparison of FBI scores between strata
  - Logistic regression interaction test: FBI x ALSFRS-R severity
    (the clinically relevant test here, distinct from the FBI x onset_site
     test already reported in 04_sensitivity_bulbar_spinal.py)

ALSFRS-R threshold justification
---------------------------------
The cut-off of ALSFRS-R <= 35 corresponds approximately to the lower
quartile of the ALSFRS-R distribution in our cohort (mean 38.7, median ~39),
thus separating patients with moderate-to-severe functional impairment from
those with mild-to-moderate impairment. This threshold was pre-specified
based on the published literature on ALSFRS-R staging (Kimura et al. 2006;
Kollewe et al. 2008) rather than derived from the data.

"""
from __future__ import annotations
from pathlib import Path
import sys
import numpy as np
import pandas as pd
from scipy.stats import mannwhitneyu
import statsmodels.api as sm
from scipy.stats import chi2

sys.path.insert(0, str(Path(__file__).parent))
from helpers import youden, perf_at, bootstrap_youden, bootstrap_ci, spawn_seeds

# ── configuration ─────────────────────────────────────────────────────────────
SEED   = 2026
N_BOOT = 2000

# ALSFRS-R threshold: lower quartile of the cohort distribution (~35),
# corresponding to moderate-severe functional impairment (see module docstring).
ALSFRS_THRESHOLD = 35

INPUT_SAV = Path("../data/cognitivita_FINAL_recoded_ETATEST_ECASFRSBEv3.sav")
INPUT_CSV = Path("../data/analytic_dataset.csv")
OUTDIR    = Path("../outputs")
OUTDIR.mkdir(parents=True, exist_ok=True)

WORD_OUTDIR = OUTDIR / "word_tables"
WORD_OUTDIR.mkdir(parents=True, exist_ok=True)


# ── data loading ──────────────────────────────────────────────────────────────
def load_data():
    """Returns the complete-case DataFrame."""
    if INPUT_SAV.exists():
        import pyreadstat
        df, _ = pyreadstat.read_sav(str(INPUT_SAV))
        sla   = df[df["SLA_NONSLA"] == "SLA"].copy()
        df    = sla[sla["FBI_total_score"].notna()].copy()
        anch  = ["ECAS_BEH_PATOL", "FRSBE_TOTALE_PATOL", "BBI_PATOL"]
        cc    = df[df[anch].notna().all(axis=1)].copy()
        cc["gold"] = (cc[anch].astype(int).sum(axis=1) >= 2).astype(int)
        return cc.rename(columns={"ALSFRSRTOT": "alsfrs_r_total"})
    elif INPUT_CSV.exists():
        cc = pd.read_csv(INPUT_CSV)
        cc = cc[cc["in_complete_case"] == 1].copy()
        cc = cc.rename(columns={
            "fbi_total"      : "FBI_total_score",
            "gold_consensus" : "gold",
        })
        return cc
    else:
        raise FileNotFoundError(
            "Provide either the SAV or the analytic CSV (paths configured at top)."
        )


# ── Word export ───────────────────────────────────────────────────────────────
def save_word_table(rows: list[dict], out_path: Path) -> None:
    """
    Generates a formatted Word document (.docx) containing Table 6
    (motor severity sensitivity analysis) using python-docx.

    Parameters
    ----------
    rows : list of dict
        Records produced by the main analysis loop, one per subgroup.
    out_path : Path
        Destination .docx file path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    # ── helpers ────────────────────────────────────────────────────────────
    def set_cell_bg(cell, hex_color: str) -> None:
        """Fill a table cell with a solid background colour."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.upper())
        tcPr.append(shd)

    def set_cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
        """Apply borders to a cell; pass a dict with 'sz' and 'color' or None."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side, spec in [("top", top), ("bottom", bottom),
                           ("left", left), ("right", right)]:
            el = OxmlElement(f"w:{side}")
            if spec:
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    str(spec["sz"]))
                el.set(qn("w:color"), spec["color"].upper())
            else:
                el.set(qn("w:val"), "none")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def set_col_width(cell, width_cm: float) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(int(width_cm * 567)))   # 1 cm ≈ 567 twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    def set_page_landscape(doc) -> None:
        """Switch the document to A4 landscape."""
        section = doc.sections[0]
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.0)
        # Mark as landscape in the XML
        sectPr = section._sectPr
        pgSz   = sectPr.find(qn("w:pgSz"))
        if pgSz is None:
            pgSz = OxmlElement("w:pgSz")
            sectPr.append(pgSz)
        pgSz.set(qn("w:orient"), "landscape")

    # ── column definitions ─────────────────────────────────────────────────
    # (header, row-dict key, format spec, column width in cm)
    columns = [
        ("Subgroup",        "subgroup",    None,    5.5),
        ("n",               "n",           "d",     1.0),
        ("Gold prev.",      "gold_prev",   ".1%",   1.5),
        ("AUC",             "auc",         ".4f",   1.4),
        ("AUC 95% CI",      "auc_ci",      None,    2.4),
        ("Youden cut",      "youden_cut",  ".0f",   1.4),
        ("Cut 95% CI",      "cut_ci",      None,    1.8),
        ("Sens @9",         "sens_at_9",   ".3f",   1.3),
        ("Spec @9",         "spec_at_9",   ".3f",   1.3),
        ("\u03ba @9",       "kappa_at_9",  "+.3f",  1.3),
        ("Sens @25",        "sens_at_25",  ".3f",   1.3),
        ("Spec @25",        "spec_at_25",  ".3f",   1.3),
    ]

    def fmt(val, spec):
        if spec is None:
            return str(val)
        if spec == "d":
            return str(int(val))
        return format(float(val), spec)

    # Pre-format rows and add combined CI columns
    fmt_rows = []
    for r in rows:
        fr = dict(r)
        fr["auc_ci"] = f"{r['auc_lo']:.4f}\u2013{r['auc_hi']:.4f}"
        fr["cut_ci"] = f"{r['cut_lo']:.0f}\u2013{r['cut_hi']:.0f}"
        fmt_rows.append([fmt(fr[c[1]], c[2]) for c in columns])

    headers   = [c[0] for c in columns]
    col_widths = [c[3] for c in columns]

    # Colours
    HEAD_BG  = "2E4057"
    HEAD_FG  = "FFFFFF"
    BODY_FG  = "000000"
    ALT_BG   = "F2F4F6"
    NOTE_FG  = "666666"

    border_med  = {"sz": "12", "color": HEAD_BG}
    border_thin = {"sz": "4",  "color": "BBBBBB"}

    # ── build document ─────────────────────────────────────────────────────
    doc = Document()
    set_page_landscape(doc)

    # Remove default paragraph spacing from Normal style
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after  = Pt(4)
    title_para.paragraph_format.space_before = Pt(0)
    run = title_para.add_run(
        "Table 6 \u2013 Sensitivity Analysis by Motor Severity (ALSFRS-R)"
    )
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # Caption
    cap_para = doc.add_paragraph()
    cap_para.paragraph_format.space_after  = Pt(8)
    cap_para.paragraph_format.space_before = Pt(0)
    cap_run = cap_para.add_run(
        "ROC recalibration of the FBI total cut-off stratified by ALSFRS-R severity "
        "(mild: ALSFRS-R\u202f>\u202f35; moderate-severe: ALSFRS-R\u202f\u2264\u202f35). "
        "AUC and Youden cut-off 95\u202f% CIs are bootstrap-based (2\u202f000 iterations). "
        "\u03ba\u202f=\u202fCohen\u2019s kappa."
    )
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Table
    n_rows = 1 + len(fmt_rows)   # header + data
    n_cols = len(columns)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for ci, (cell, hdr, width) in enumerate(zip(hdr_row.cells, headers, col_widths)):
        set_col_width(cell, width)
        set_cell_bg(cell, HEAD_BG)
        set_cell_borders(cell, top=border_med, bottom=border_med)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(hdr)
        run.bold      = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row_vals in enumerate(fmt_rows):
        tbl_row = tbl.rows[ri + 1]
        bg = ALT_BG if ri % 2 == 0 else "FFFFFF"
        for ci, (cell, val, width) in enumerate(zip(tbl_row.cells, row_vals, col_widths)):
            set_col_width(cell, width)
            set_cell_bg(cell, bg)
            set_cell_borders(cell, top=border_thin, bottom=border_thin)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(val)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Footnote
    foot_para = doc.add_paragraph()
    foot_para.paragraph_format.space_before = Pt(6)
    foot_para.paragraph_format.space_after  = Pt(0)
    foot_run = foot_para.add_run(
        "Abbreviations: AUC, area under the ROC curve; CI, confidence interval; "
        "Sens, sensitivity; Spec, specificity; \u03ba, Cohen\u2019s kappa; "
        "@9\u202f=\u202fFBI\u202f\u2265\u202f9; @25\u202f=\u202fFBI\u202f\u2265\u202f25."
    )
    foot_run.font.size = Pt(8)
    foot_run.font.italic = True
    foot_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(str(out_path))
    print(f"Saved Word table: {out_path}")


# ── main analysis ─────────────────────────────────────────────────────────────
def main():
    np.random.seed(SEED)
    cc = load_data()

    print("=" * 78)
    print("SENSITIVITY ANALYSIS BY MOTOR SEVERITY (ALSFRS-R)")
    print(f"Threshold: ALSFRS-R > {ALSFRS_THRESHOLD} (mild) vs <= {ALSFRS_THRESHOLD} (moderate-severe)")
    print("=" * 78)

    mask = cc["alsfrs_r_total"].notna()
    sub  = cc[mask].copy()
    print(f"\nComplete-case with ALSFRS-R available: n = {len(sub)} of {len(cc)}")
    print(f"ALSFRS-R: median {sub['alsfrs_r_total'].median():.0f}, "
          f"mean {sub['alsfrs_r_total'].mean():.1f} "
          f"+- {sub['alsfrs_r_total'].std():.1f}")

    sub["severity"] = np.where(
        sub["alsfrs_r_total"] > ALSFRS_THRESHOLD,
        f"Mild (ALSFRS-R >{ALSFRS_THRESHOLD})",
        f"Moderate-severe (ALSFRS-R <={ALSFRS_THRESHOLD})",
    )

    # Independent seeds for each subgroup bootstrap
    seed_mild, seed_severe = spawn_seeds(SEED, 2)

    rows = []
    for label, bseed in [
        (f"Mild (ALSFRS-R >{ALSFRS_THRESHOLD})",         seed_mild),
        (f"Moderate-severe (ALSFRS-R <={ALSFRS_THRESHOLD})", seed_severe),
    ]:
        X = sub[sub["severity"] == label]
        if len(X) < 20:
            print(f"\n{label}: n={len(X)} too small (< 20), skipping.")
            continue
        y = X["gold"].astype(int).values
        s = X["FBI_total_score"].values
        r = youden(y, s)
        cb, ab   = bootstrap_youden(y, s, n_boot=N_BOOT, seed=bseed, stratified=True)
        lo_cut, hi_cut = bootstrap_ci(cb)
        lo_auc, hi_auc = bootstrap_ci(ab)
        p9  = perf_at(y, s, 9)
        p25 = perf_at(y, s, 25)

        print(f"\n--- {label}  (n = {len(X)}, gold prev = {y.mean():.1%}) ---")
        print(f"  ALSFRS-R range : [{X['alsfrs_r_total'].min():.0f}, {X['alsfrs_r_total'].max():.0f}]")
        print(f"  FBI median     : {X['FBI_total_score'].median():.0f}")
        print(f"  AUC = {r['auc']:.4f}  (95% CI {lo_auc:.4f}-{hi_auc:.4f})")
        print(f"  Youden cut-off = FBI >= {r['cut']:.0f}  (95% CI {lo_cut:.0f}-{hi_cut:.0f})")
        print(f"  At FBI >= 9 : sens={p9['sens']:.3f}, spec={p9['spec']:.3f}, "
              f"kappa={p9['kappa']:+.3f}")
        print(f"  At FBI >= 25: sens={p25['sens']:.3f}, spec={p25['spec']:.3f}, "
              f"kappa={p25['kappa']:+.3f}")

        rows.append(dict(
            subgroup=label, n=len(X), gold_prev=y.mean(),
            auc=r["auc"], auc_lo=lo_auc, auc_hi=hi_auc,
            youden_cut=r["cut"], cut_lo=lo_cut, cut_hi=hi_cut,
            sens_at_9=p9["sens"],  spec_at_9=p9["spec"],  kappa_at_9=p9["kappa"],
            sens_at_25=p25["sens"], spec_at_25=p25["spec"],
        ))

    # FBI score comparison between severity strata
    mild  = sub.loc[sub["severity"].str.startswith("Mild"), "FBI_total_score"]
    sev   = sub.loc[sub["severity"].str.startswith("Moderate"), "FBI_total_score"]
    u_res = mannwhitneyu(mild, sev)
    print(f"\nMann-Whitney FBI mild vs moderate-severe: p = {u_res.pvalue:.4f}")
    print(f"  Mild median {mild.median():.0f}, moderate-severe median {sev.median():.0f}")
    print("  NOTE: this tests raw score location, NOT the FBI-impairment "
          "relationship; the interaction test below addresses that directly.")

    # ── FBI x ALSFRS-R severity INTERACTION TEST ───────────────────────────
    print("\n" + "=" * 78)
    print("INTERACTION TEST  FBI x ALSFRS-R severity  (logistic regression)")
    print("=" * 78)
    print("Does the FBI-impairment relationship itself differ between mild")
    print("and moderate-severe patients? This is the clinically relevant")
    print("test for potential motor confound via ALSFRS-R.")

    Xint = sub[["FBI_total_score", "severity", "gold"]].dropna().copy()
    Xint["severe"]      = Xint["severity"].str.startswith("Moderate").astype(int)
    Xint["fbi"]         = Xint["FBI_total_score"].astype(float)
    Xint["fbi_x_sev"]   = Xint["fbi"] * Xint["severe"]
    design = sm.add_constant(Xint[["fbi", "severe", "fbi_x_sev"]])
    p_int = None
    try:
        mod_int = sm.Logit(Xint["gold"].astype(int), design).fit(disp=0)
        print(mod_int.summary().as_text())
        p_int = float(mod_int.pvalues["fbi_x_sev"])
    except np.linalg.LinAlgError:
        print("  WARNING: unpenalised Logit did not converge (singular Hessian). "
              "Refitting with L2 penalty (alpha=0.1) -- p-value is approximate.")
        mod_int   = sm.Logit(Xint["gold"].astype(int), design).fit_regularized(
            alpha=0.1, disp=0)
        X_noint   = sm.add_constant(Xint[["fbi", "severe"]])
        mod_noint = sm.Logit(Xint["gold"].astype(int), X_noint).fit_regularized(
            alpha=0.1, disp=0)
        lr    = max(2 * (mod_int.llf - mod_noint.llf), 0)
        p_int = float(1 - chi2.cdf(lr, df=1))
        print(f"  Penalised LR test: LR={lr:.3f}, approx p={p_int:.4f}")

    print(f"\nFBI x severity interaction p-value = {p_int:.4f}")
    if p_int < 0.05:
        print("  WARNING: significant interaction -- the FBI-impairment "
              "slope differs between mild and moderate-severe patients. "
              "A single cut-off may not be fully appropriate across severity strata.")
    else:
        print("  Non-significant interaction -- consistent with a shared "
              "FBI-impairment relationship across ALSFRS-R strata.")

    # ── save outputs ──────────────────────────────────────────────────────────
    csv_path = OUTDIR / "table6_motor_severity.csv"
    pd.DataFrame(rows).to_csv(csv_path, index=False)
    print(f"\nSaved CSV: {csv_path}")

    if rows:
        word_path = WORD_OUTDIR / "table6_motor_severity.docx"
        save_word_table(rows, word_path)
    else:
        print("No subgroups with sufficient n — Word table not generated.")


if __name__ == "__main__":
    main()